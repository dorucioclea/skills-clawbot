#!/usr/bin/env python3
"""
TubeScribe - YouTube Video Summarizer
=====================================

Extracts transcripts from YouTube videos, generates summaries with speaker
detection, and creates document + audio output.

Usage:
    python tubescribe.py <youtube_url> [options]

Options:
    --output-dir DIR     Output directory (default: current dir)
    --doc-format FORMAT  Document format: rtf, docx, md (default: from config)
    --audio-format FMT   Audio format: wav, mp3, none (default: from config)
    --no-audio           Skip audio generation
    --quiet              Minimal output

Author: Jackie 🦊 & Matus
Created: 2026-02-04
"""

import subprocess
import json
import re
import sys
import os
import urllib.request
from pathlib import Path

# Add scripts dir to path for imports
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR))

from rtf_writer import create_rtf_from_markdown

# Config
CONFIG_FILE = os.path.expanduser("~/.tubescribe/config.json")
DEFAULT_CONFIG = {
    "document_format": "rtf",
    "audio_format": "wav",
    "tts_engine": "builtin",
}


def load_config() -> dict:
    """Load config or return defaults."""
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r') as f:
            return {**DEFAULT_CONFIG, **json.load(f)}
    return DEFAULT_CONFIG


def extract_video_id(url: str) -> str:
    """Extract video ID from various YouTube URL formats."""
    patterns = [
        r'(?:v=|/v/|youtu\.be/)([a-zA-Z0-9_-]{11})',
        r'(?:embed/)([a-zA-Z0-9_-]{11})',
        r'^([a-zA-Z0-9_-]{11})$'
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return "unknown"


def get_video_metadata(url: str) -> dict:
    """Extract video title and description from YouTube."""
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=10) as response:
            html = response.read().decode('utf-8')
        
        title_match = re.search(r'"title":"([^"]+)"', html)
        title = title_match.group(1) if title_match else "Unknown"
        title = title.encode().decode('unicode_escape')
        
        desc_match = re.search(r'"shortDescription":"([^"]*)"', html)
        description = desc_match.group(1) if desc_match else ""
        description = description.encode().decode('unicode_escape')
        
        return {"title": title, "description": description}
    except Exception as e:
        print(f"   Warning: Could not fetch metadata: {e}")
        return {"title": "Unknown Video", "description": ""}


def download_transcript(url: str) -> list[dict]:
    """Download transcript with timestamps using summarize CLI."""
    try:
        result = subprocess.run(
            ["summarize", url, "--youtube", "auto", "--extract-only", "--timestamps"],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            raise Exception(f"summarize failed: {result.stderr}")
        
        lines = result.stdout.strip().split('\n')
        segments = []
        
        for line in lines:
            match = re.match(r'\[(\d+:\d+(?::\d+)?)\]\s*(.*)', line)
            if match:
                segments.append({
                    "ts": match.group(1),
                    "text": match.group(2).strip()
                })
        
        return segments
    except FileNotFoundError:
        print("❌ Error: 'summarize' CLI not found. Install with: brew install steipete/tap/summarize")
        sys.exit(1)


def clean_transcript(segments: list[dict]) -> list[dict]:
    """Clean transcript segments."""
    cleaned = []
    for seg in segments:
        text = seg["text"]
        text = re.sub(r'\s+', ' ', text).strip()
        text = re.sub(r'>>\s*', '', text)  # Remove >> artifacts
        if text:
            cleaned.append({"ts": seg["ts"], "text": text})
    return cleaned


def sanitize_filename(title: str) -> str:
    """Create safe filename from title."""
    safe = re.sub(r'[<>:"/\\|?*]', '', title)
    safe = re.sub(r'\s+', ' ', safe).strip()
    return safe[:100]  # Limit length


def prepare_source_data(url: str, output_dir: str = "/tmp") -> tuple[str, str]:
    """
    Prepare source data for processing.
    Returns: (source_json_path, output_md_path)
    """
    video_id = extract_video_id(url)
    
    print(f"🎬 TubeScribe")
    print(f"   Video ID: {video_id}")
    
    print("1. Fetching video metadata...")
    metadata = get_video_metadata(url)
    print(f"   Title: {metadata['title'][:60]}...")
    
    print("2. Downloading transcript...")
    segments = download_transcript(url)
    print(f"   Raw segments: {len(segments)}")
    
    print("3. Cleaning transcript...")
    segments = clean_transcript(segments)
    word_count = sum(len(s["text"].split()) for s in segments)
    print(f"   Segments: {len(segments)}")
    print(f"   Words: {word_count}")
    
    # Prepare output
    source_data = {
        "metadata": {
            "url": url,
            "title": metadata["title"],
            "description": metadata["description"],
            "video_id": video_id,
        },
        "segments": segments
    }
    
    source_path = f"{output_dir}/tubescribe_{video_id}_source.json"
    output_path = f"{output_dir}/tubescribe_{video_id}_output.md"
    
    with open(source_path, 'w', encoding='utf-8') as f:
        json.dump(source_data, f, indent=2, ensure_ascii=False)
    
    print(f"\n✅ Source: {source_path}")
    print(f"📝 Output: {output_path}")
    
    return source_path, output_path


def convert_to_document(md_path: str, output_dir: str, doc_format: str) -> str:
    """Convert markdown to final document format."""
    
    # Get title from markdown for filename
    with open(md_path, 'r') as f:
        content = f.read()
    title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
    title = title_match.group(1) if title_match else "TubeScribe Output"
    safe_title = sanitize_filename(title)
    
    if doc_format == "md":
        # Just copy/rename markdown
        out_path = os.path.join(output_dir, f"{safe_title}.md")
        with open(out_path, 'w') as f:
            f.write(content)
        return out_path
    
    elif doc_format == "rtf":
        out_path = os.path.join(output_dir, f"{safe_title}.rtf")
        create_rtf_from_markdown(md_path, out_path)
        return out_path
    
    elif doc_format == "docx":
        out_path = os.path.join(output_dir, f"{safe_title}.docx")
        # Try pandoc first
        try:
            subprocess.run(
                ["pandoc", md_path, "-o", out_path],
                check=True, capture_output=True
            )
            return out_path
        except FileNotFoundError:
            pass
        
        # Try python-docx
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            
            doc = Document()
            # Simple conversion - headers and paragraphs
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    # Handle bold
                    p = doc.add_paragraph()
                    # Simple bold handling
                    parts = re.split(r'\*\*(.+?)\*\*', line)
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold content
                            run.bold = True
            
            doc.save(out_path)
            return out_path
        except ImportError:
            print("   Warning: Neither pandoc nor python-docx available, falling back to RTF")
            return convert_to_document(md_path, output_dir, "rtf")
    
    return md_path


def generate_audio_summary(summary_text: str, output_path: str, config: dict) -> str:
    """Generate audio summary using configured TTS engine."""
    
    audio_format = config.get("audio_format", "wav")
    tts_engine = config.get("tts_engine", "builtin")
    
    print(f"🔊 Generating audio ({tts_engine}, {audio_format})...")
    
    if tts_engine == "kokoro":
        return generate_kokoro_audio(summary_text, output_path, audio_format)
    else:
        return generate_builtin_audio(summary_text, output_path, audio_format)


def get_kokoro_python() -> str:
    """Get the Python interpreter for Kokoro (venv or system)."""
    config = load_config()
    venv_path = config.get("kokoro_venv", os.path.expanduser("~/.tubescribe/kokoro-env"))
    venv_python = os.path.join(venv_path, "bin", "python")
    
    if os.path.exists(venv_python):
        return venv_python
    return sys.executable


def generate_kokoro_audio(text: str, output_path: str, audio_format: str) -> str:
    """Generate audio using Kokoro TTS."""
    try:
        wav_path = output_path.replace('.mp3', '.wav') if output_path.endswith('.mp3') else output_path
        if not wav_path.endswith('.wav'):
            wav_path = output_path + '.wav'
        
        kokoro_python = get_kokoro_python()
        
        # Escape text for Python string
        escaped_text = text.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')
        
        code = f'''
from kokoro import KPipeline
import soundfile as sf
import torch
import numpy as np
import os

pipeline = KPipeline(lang_code='a')

# Try to load custom voice blend, fall back to default
try:
    cache_dir = os.path.expanduser('~/.cache/huggingface/hub/models--hexgrad--Kokoro-82M/snapshots')
    snapshot = os.listdir(cache_dir)[0]
    voices_dir = os.path.join(cache_dir, snapshot, 'voices')
    heart = torch.load(f'{{voices_dir}}/af_heart.pt', weights_only=True)
    sky = torch.load(f'{{voices_dir}}/af_sky.pt', weights_only=True)
    voice = 0.6 * heart + 0.4 * sky
except:
    voice = 'af_heart'

text = "{escaped_text}"
audio_chunks = []
for _, _, audio in pipeline(text, voice=voice):
    audio_chunks.append(audio)
full_audio = np.concatenate(audio_chunks)
sf.write("{wav_path}", full_audio, 24000)
print("OK")
'''
        result = subprocess.run([kokoro_python, "-c", code], capture_output=True, text=True, timeout=300)
        
        if result.returncode == 0 and os.path.exists(wav_path):
            if audio_format == "mp3":
                mp3_path = output_path if output_path.endswith('.mp3') else output_path.replace('.wav', '.mp3')
                subprocess.run(["ffmpeg", "-y", "-i", wav_path, "-b:a", "128k", mp3_path], 
                             capture_output=True, check=True)
                os.remove(wav_path)
                return mp3_path
            return wav_path
        else:
            raise Exception(f"Kokoro failed: {result.stderr}")
    
    except Exception as e:
        print(f"   Kokoro failed, falling back to built-in TTS: {e}")
        return generate_builtin_audio(text, output_path, audio_format)


def generate_builtin_audio(text: str, output_path: str, audio_format: str) -> str:
    """Generate audio using macOS say command (fallback)."""
    wav_path = output_path.replace('.mp3', '.wav')
    aiff_path = output_path.replace('.mp3', '.aiff').replace('.wav', '.aiff')
    
    # Use macOS say command
    subprocess.run(["say", "-o", aiff_path, text], check=True, capture_output=True)
    
    # Convert to wav
    subprocess.run(["afconvert", "-f", "WAVE", "-d", "LEI16", aiff_path, wav_path],
                  check=True, capture_output=True)
    os.remove(aiff_path)
    
    if audio_format == "mp3":
        mp3_path = output_path
        try:
            subprocess.run(["ffmpeg", "-y", "-i", wav_path, "-b:a", "128k", mp3_path],
                         capture_output=True, check=True)
            os.remove(wav_path)
            return mp3_path
        except FileNotFoundError:
            print("   ffmpeg not found, keeping wav format")
            return wav_path
    
    return wav_path


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description="TubeScribe - YouTube Video Summarizer")
    parser.add_argument("url", help="YouTube URL")
    parser.add_argument("--output-dir", default=".", help="Output directory")
    parser.add_argument("--doc-format", choices=["rtf", "docx", "md"], help="Document format")
    parser.add_argument("--audio-format", choices=["wav", "mp3", "none"], help="Audio format")
    parser.add_argument("--no-audio", action="store_true", help="Skip audio generation")
    parser.add_argument("--quiet", action="store_true", help="Minimal output")
    
    args = parser.parse_args()
    
    config = load_config()
    
    # Override config with args
    if args.doc_format:
        config["document_format"] = args.doc_format
    if args.audio_format:
        config["audio_format"] = args.audio_format
    
    # Prepare source data
    source_path, output_md_path = prepare_source_data(args.url)
    
    print("\n📋 Next step: Process with AI agent")
    print(f"   Source: {source_path}")
    print(f"   Output: {output_md_path}")
    print("\nThe agent will:")
    print("  1. Read the source JSON")
    print("  2. Detect speakers and create summary")
    print("  3. Write formatted markdown to output path")
    print("\nThen run this script again with --finalize to create documents.")


if __name__ == "__main__":
    main()
